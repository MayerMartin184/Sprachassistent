"""Dateien auf dem Rechner: suchen, lesen, ablegen, ordnen, öffnen – in freigegebenen Wurzelordnern.

Pfade werden mit dem Namen der Wurzel angegeben, z. B. "Downloads/Rechnung.pdf" oder "Dokumente/Projekte".
Außerhalb der Wurzeln passiert nichts. Löschen geht in den Papierkorb und braucht eine Bestätigung.
"""

from __future__ import annotations

import logging
import os
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Callable

from .base import Tool, schema

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log", ".ini", ".xml", ".html", ".py", ".rtf"}
DOC_SUFFIXES = {".pdf", ".docx", ".xlsx"}
MAX_RESULTS = 60
SKIP_DIRS = {"node_modules", ".git", "__pycache__", "AppData", "$RECYCLE.BIN", "System Volume Information"}


def default_roots() -> dict[str, Path]:
    home = Path.home()
    roots = {"Dokumente": home / "Documents", "Downloads": home / "Downloads", "Desktop": home / "Desktop"}
    onedrive = os.environ.get("OneDrive")
    if onedrive:
        roots["OneDrive"] = Path(onedrive)
    return {name: path for name, path in roots.items() if path.exists()}


def parse_roots(spec: str | None) -> dict[str, Path]:
    """FILE_ROOTS="Projekte=D:\\Projekte;Server=\\\\srv\\daten" -> {"Projekte": Path(...), ...}"""
    roots: dict[str, Path] = {}
    for part in (spec or "").split(";"):
        if "=" in part:
            name, _, path = part.partition("=")
            if name.strip() and path.strip():
                roots[name.strip()] = Path(path.strip()).expanduser()
    return roots


class FileManager:
    def __init__(self, roots: dict[str, Path], confirm: Callable[[str], bool] | None = None) -> None:
        self.roots = {name: path.expanduser().resolve() for name, path in roots.items()}
        for path in self.roots.values():
            path.mkdir(parents=True, exist_ok=True)
        self.confirm = confirm or (lambda _m: True)

    # --- Pfade -----------------------------------------------------------------
    def resolve(self, spec: str) -> Path:
        spec = (spec or "").strip().replace("\\", "/")
        if not spec or spec == ".":
            raise ValueError("Bitte Wurzel angeben, z. B. 'Dokumente' oder 'Downloads/Datei.pdf'. Wurzeln: " + ", ".join(self.roots))
        candidate = Path(spec)
        if candidate.is_absolute():
            path = candidate.resolve()
        else:
            root_name, _, rest = spec.partition("/")
            root = self._root(root_name)
            path = (root / rest).resolve() if rest else root
        for root in self.roots.values():
            if path == root or root in path.parents:
                return path
        raise PermissionError(f"Pfad liegt außerhalb der freigegebenen Ordner ({', '.join(self.roots)}): {spec}")

    def _root(self, name: str) -> Path:
        for key, path in self.roots.items():
            if key.lower() == name.lower():
                return path
        raise KeyError(f"Unbekannte Wurzel '{name}'. Verfügbar: " + ", ".join(self.roots))

    def rel(self, path: Path) -> str:
        for name, root in self.roots.items():
            if path == root:
                return name
            if root in path.parents:
                return f"{name}/{path.relative_to(root).as_posix()}"
        return str(path)

    # --- Lesen -----------------------------------------------------------------
    def list_dir(self, spec: str) -> str:
        path = self.resolve(spec)
        if not path.is_dir():
            raise FileNotFoundError(f"Ordner nicht gefunden: {spec}")
        entries = sorted(path.iterdir(), key=lambda p: (p.is_file(), p.name.lower()))
        if not entries:
            return f"Ordner '{self.rel(path)}' ist leer."
        lines = [f"Inhalt von '{self.rel(path)}':"] + [self._describe(e) for e in entries[:MAX_RESULTS]]
        if len(entries) > MAX_RESULTS:
            lines.append(f"... und {len(entries) - MAX_RESULTS} weitere")
        return "\n".join(lines)

    def search(self, query: str, spec: str | None = None, content: bool = False) -> str:
        bases = [self.resolve(spec)] if spec else list(self.roots.values())
        needle = query.lower()
        hits: list[Path] = []
        for base in bases:
            for dirpath, dirnames, filenames in os.walk(base):
                dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS and not d.startswith(".")]
                for name in dirnames + filenames:
                    if needle in name.lower():
                        hits.append(Path(dirpath) / name)
                if content:
                    for name in filenames:
                        p = Path(dirpath) / name
                        if p.suffix.lower() in TEXT_SUFFIXES | DOC_SUFFIXES and p not in hits and needle in self._text_of(p, 200000).lower():
                            hits.append(p)
                if len(hits) >= MAX_RESULTS * 3:
                    break
        if not hits:
            return f"Keine Treffer für '{query}'."
        hits.sort(key=lambda p: p.stat().st_mtime if p.exists() else 0, reverse=True)
        return f"{len(hits)} Treffer für '{query}':\n" + "\n".join(self._describe(p) for p in hits[:MAX_RESULTS])

    def read(self, spec: str, max_chars: int = 8000) -> str:
        path = self.resolve(spec)
        if not path.is_file():
            raise FileNotFoundError(f"Datei nicht gefunden: {spec}")
        text = self._text_of(path, max_chars * 4)
        if len(text) > max_chars:
            return text[:max_chars] + f"\n... [gekürzt, {len(text)} Zeichen gesamt]"
        return text or "(kein Text gefunden)"

    def _text_of(self, path: Path, limit: int) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix in TEXT_SUFFIXES:
                return path.read_text(encoding="utf-8", errors="replace")[:limit]
            if suffix == ".pdf":
                from pypdf import PdfReader

                parts = []
                for page in PdfReader(str(path)).pages:
                    parts.append(page.extract_text() or "")
                    if sum(len(x) for x in parts) > limit:
                        break
                return "\n".join(parts)[:limit]
            if suffix == ".docx":
                import docx

                d = docx.Document(str(path))
                parts = [p.text for p in d.paragraphs]
                for table in d.tables:
                    for row in table.rows:
                        parts.append(" | ".join(c.text for c in row.cells))
                return "\n".join(parts)[:limit]
            if suffix == ".xlsx":
                import openpyxl

                wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                parts = []
                for ws in wb.worksheets:
                    parts.append(f"## Blatt {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        cells = ["" if v is None else str(v) for v in row]
                        if any(cells):
                            parts.append(" | ".join(cells))
                        if sum(len(x) for x in parts) > limit:
                            break
                return "\n".join(parts)[:limit]
        except ImportError as exc:
            raise ValueError(f"Zum Lesen von {suffix} fehlt ein Paket: {exc.name}") from exc
        raise ValueError(f"Kein lesbares Format: {suffix}")

    # --- Schreiben / Ordnen -------------------------------------------------------
    def mkdir(self, spec: str) -> str:
        path = self.resolve(spec)
        path.mkdir(parents=True, exist_ok=True)
        return f"Ordner vorhanden: {self.rel(path)}"

    def write(self, spec: str, content: str, overwrite: bool = False) -> str:
        path = self.resolve(spec)
        if path.exists() and not overwrite:
            raise FileExistsError(f"Datei existiert bereits (overwrite=true zum Überschreiben): {spec}")
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return f"Gespeichert: {self.rel(path)} ({len(content)} Zeichen)"

    def _target(self, src: Path, destination: str) -> Path:
        dst = self.resolve(destination)
        if dst.is_dir():
            dst = dst / src.name
        if dst.exists():
            raise FileExistsError(f"Ziel existiert bereits: {self.rel(dst)}")
        dst.parent.mkdir(parents=True, exist_ok=True)
        return dst

    def move(self, source: str, destination: str) -> str:
        src = self.resolve(source)
        if not src.exists():
            raise FileNotFoundError(f"Quelle nicht gefunden: {source}")
        dst = self._target(src, destination)
        shutil.move(str(src), str(dst))
        return f"Verschoben: {self.rel(src)} -> {self.rel(dst)}"

    def copy(self, source: str, destination: str) -> str:
        src = self.resolve(source)
        if not src.exists():
            raise FileNotFoundError(f"Quelle nicht gefunden: {source}")
        dst = self._target(src, destination)
        if src.is_dir():
            shutil.copytree(src, dst)
        else:
            shutil.copy2(src, dst)
        return f"Kopiert: {self.rel(src)} -> {self.rel(dst)}"

    def delete(self, spec: str) -> str:
        path = self.resolve(spec)
        if path in self.roots.values():
            raise PermissionError("Wurzelordner werden nicht gelöscht.")
        if not path.exists():
            raise FileNotFoundError(f"Nicht gefunden: {spec}")
        kind = "Ordner" if path.is_dir() else "Datei"
        if not self.confirm(f"{kind} in den Papierkorb verschieben?\n{self.rel(path)}"):
            return "Der Nutzer hat das Löschen abgelehnt."
        try:
            from send2trash import send2trash

            send2trash(str(path))
            return f"In den Papierkorb verschoben: {self.rel(path)}"
        except ImportError:
            trash = Path.home() / ".sprachassistent" / "papierkorb"
            trash.mkdir(parents=True, exist_ok=True)
            target = trash / f"{datetime.now():%Y%m%d_%H%M%S}_{path.name}"
            shutil.move(str(path), str(target))
            return f"Verschoben nach {target} (Papierkorb-Paket fehlt)."

    def open(self, spec: str) -> str:
        path = self.resolve(spec)
        if not path.exists():
            raise FileNotFoundError(f"Nicht gefunden: {spec}")
        if sys.platform == "win32":
            os.startfile(str(path))  # type: ignore[attr-defined]
        else:
            import subprocess

            subprocess.Popen(["open" if sys.platform == "darwin" else "xdg-open", str(path)])
        return f"Geöffnet: {self.rel(path)}"

    def _describe(self, path: Path) -> str:
        try:
            stat = path.stat()
        except OSError:
            return f"[?]      {self.rel(path)}"
        mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M")
        if path.is_dir():
            return f"[Ordner] {self.rel(path)}/"
        return f"[Datei]  {self.rel(path)}  ({_human(stat.st_size)}, {mtime})"


def _human(size: float) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if size < 1024 or unit == "GB":
            return f"{size:.0f} {unit}" if unit == "B" else f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} GB"


def build_tools(fm: FileManager) -> list[Tool]:
    roots = ", ".join(f"{n} ({p})" for n, p in fm.roots.items())
    hint = f"Pfade beginnen mit dem Namen der Wurzel, z. B. 'Downloads/Rechnung.pdf'. Wurzeln: {roots}."
    return [
        Tool("files_list", f"Listet den Inhalt eines Ordners. {hint}", schema({"path": {"type": "string"}}, ["path"]), lambda path: fm.list_dir(path)),
        Tool(
            "files_search",
            f"Sucht Dateien/Ordner per Namensbestandteil in allen Wurzeln oder unter path; content=true sucht zusätzlich im Text von txt/pdf/docx/xlsx (langsamer). {hint}",
            schema({"query": {"type": "string"}, "path": {"type": "string"}, "content": {"type": "boolean"}}, ["query"]),
            lambda query, path=None, content=False: fm.search(query, path, content),
        ),
        Tool(
            "files_read",
            f"Liest den Text einer Datei: txt, md, csv, json, pdf, docx, xlsx. {hint}",
            schema({"path": {"type": "string"}, "max_chars": {"type": "integer", "description": "Standard 8000"}}, ["path"]),
            lambda path, max_chars=8000: fm.read(path, max_chars),
        ),
        Tool("files_write", f"Schreibt eine Textdatei (Notizen, Rechercheergebnisse, Reviews). {hint}",
             schema({"path": {"type": "string"}, "content": {"type": "string"}, "overwrite": {"type": "boolean"}}, ["path", "content"]),
             lambda path, content, overwrite=False: fm.write(path, content, overwrite)),
        Tool("files_mkdir", f"Legt einen Ordner an. {hint}", schema({"path": {"type": "string"}}, ["path"]), lambda path: fm.mkdir(path)),
        Tool("files_move", f"Verschiebt oder benennt um; ist das Ziel ein Ordner, wird dort abgelegt. Überschreibt nie. {hint}",
             schema({"source": {"type": "string"}, "destination": {"type": "string"}}, ["source", "destination"]), fm.move),
        Tool("files_copy", f"Kopiert Datei oder Ordner. {hint}",
             schema({"source": {"type": "string"}, "destination": {"type": "string"}}, ["source", "destination"]), fm.copy),
        Tool("files_delete", f"Verschiebt Datei oder Ordner in den Papierkorb; der Nutzer muss bestätigen. {hint}",
             schema({"path": {"type": "string"}}, ["path"]), lambda path: fm.delete(path)),
        Tool("files_open", f"Öffnet Datei oder Ordner mit dem Standardprogramm (z. B. PDF im Reader, Ordner im Explorer). {hint}",
             schema({"path": {"type": "string"}}, ["path"]), lambda path: fm.open(path)),
    ]
